import pandas as pd
from pybaseball import statcast_pitcher
from pybaseball import playerid_lookup
import matplotlib.pylab as plt
import matplotlib.gridspec as gridspec
import math as math
import datetime
from pandas.compat import BytesIO
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
import numpy as np
import os
from plotly.subplots import make_subplots
import plotly.graph_objects as go
from chart_studio.plotly import image as PlotlyImage
import psutil
import plotly.io as pio
import plotly.express as px
import requests
import io
from typing import Optional, Union
import warnings


plt.style.use('seaborn-paper')


# used for debugging
pd.set_option('display.max_columns', None)
pd.set_option('display.expand_frame_repr', False)
pd.set_option('max_colwidth', -1)


def get_number(last, first, year):
    playerTable = playerid_lookup(last, first)
    playerTable = playerTable.loc[playerTable['mlb_played_last'].isin(['2020','2021'])]
    playerTable.index = range(len(playerTable['mlb_played_last']))
    number = playerTable['key_mlbam']
    number = number[0]
    return number


def statcast_pitcher_active_spin(year: int, minP: int = 1, _type: str = 'spin-based', player_id: Optional[int] = None) -> pd.DataFrame:
    # Statcast supports spin-based for some years, but not others. We'll try to get that first, but if it's empty
    # we'll fall back to the observed.
    #
    # From Statcast:
    #   Measured active spin uses the 3D spin vector at release; this is only possible with the 2020 season going
    #   forward. (Label is "2020 - Spin-based" and can be read as "Active Spin using the Spin-based method".)
    #   Inferred active spin from movement uses the total amount of movement to estimate the amount of active spin,
    #   if we presumed only magnus was at play; this is a legacy method that can be useful in certain circumstances.
    #   (Label is "2020 - Observed" and can be read as "Active Spin using the Total Observed Movement method".)


    #url = f"https://baseballsavant.mlb.com/leaderboard/active-spin?year={year}_{_type}&min={minP}&hand=&csv=true"
    url = f"https://baseballsavant.mlb.com/leaderboard/spin-direction-pitches?year={year}&min={minP}&sort=0&sortDir=desc&pitch_type=ALL&throws=&playerName=&team=&pov=Pit&csv=true"
    res = requests.get(url, timeout=None).content
    if res and '<html' in res.decode('utf-8'):
        # This did no go as planned. Statcast redirected us back to HTML :(
        if _type == 'spin-based':
            warnings.warn(f'Could not get active spin results for year {year} that are "spin-based". Trying to get the older "observed" results.')
            return statcast_pitcher_active_spin(year, minP, 'observed')

        warnings.warn("Statcast did not return any active spin results for the query provided.")
        return pd.DataFrame()
    data = pd.read_csv(io.StringIO(res.decode('utf-8')))
    if _type == 'spin-based' and (data is None or data.empty):
        return statcast_pitcher_active_spin(year, minP, 'observed')

    if player_id != np.nan:
            data = data.loc[data['player_id'] == player_id]

    return data


def import_data(number, start, end, year):
    data = statcast_pitcher(start_dt=start, end_dt=end,
                            player_id=number)
    data = data[['pitch_type', 'release_speed', 'release_pos_x', 'release_pos_z',
                 'pfx_x', 'pfx_z', 'release_spin_rate', 'plate_x', 'plate_z',
                 'estimated_woba_using_speedangle', 'woba_value', 'description',
                 'launch_speed_angle', 'launch_angle', 'launch_speed', 'bb_type',
                 'effective_speed', 'vx0', 'vy0', 'vz0', 'ax', 'ay', 'az',
                 'release_extension', 'sz_top', 'sz_bot', 'p_throws', 'spin_axis',
                 'delta_run_exp']]
    data.index = range(len(data['pitch_type']))

    spin_data = statcast_pitcher_active_spin(year = year, player_id = number)

    spin_data = spin_data.rename(columns={"api_pitch_type": "pitch_type"})
    spin_data = spin_data.drop(columns = ['year', 'last_name', ' first_name',
                                          'pitch_hand', 'n_pitches', 'release_speed',
                                          'spin_rate', 'movement_inches', 'hawkeye_measured',
                                          'movement_inferred', 'api_pitch_name', 'active_spin_formatted',
                                          'hawkeye_measured_clock_minutes', 'movement_inferred_clock_minutes',
                                          'diff_measured_inferred', 'diff2','diff_measured_inferred_minutes',
                                          'hawkeye_measured_clock_hh', 'hawkeye_measured_clock_mm',
                                          'movement_inferred_clock_hh', 'movement_inferred_clock_mm',
                                          'diff_clock_hh', 'diff_clock_mm', 'hawkeye_measured_clock_label',
                                          'movement_inferred_clock_label', 'diff_clock_label'])
    return data, spin_data


def nathan_calculations(pitches):
    # get avg weather and constants
    weather_data = pd.read_csv('AvgWeather.csv')
    #player_team = weather_data.loc[weather_data['team'].isin([team])]
    #home_temp = player_team.temp.mean()
    #home_pressure = player_team.adj_pressure.mean()
    #home_humidity = player_team.humidity.mean()

    #other = weather_data.loc[~weather_data['team'].isin([team])]
    other_temp = weather_data.temp.mean()
    other_pressure = weather_data.adj_pressure.mean()
    other_humidity = weather_data.humidity.mean()

    g_fts = 32.174
    R_ball = .121
    mass = 5.125
    circ = 9.125
    temp = other_temp
    humidity = other_humidity
    pressure = other_pressure
    temp_c = (5/9)*(temp-32)
    pressure_mm = (pressure * 1000) / 39.37
    svp = 4.5841 * math.exp((18.687 - temp_c/234.5) * temp_c/(257.14 + temp_c))
    rho = (1.2929 * (273 / (temp_c + 273)) * (pressure_mm - .3783 * humidity * svp / 100) / 760) * .06261
    #const = 0.07182 * rho * (5.125 / mass) * (circ / 9.125)**2
    const = 0.00538309893013877
    # add row to put calculations in
    pitches['InducedHorzBreak'] = np.nan
    pitches['total_mov'] = np.nan
    pitches['InducedVertBreak'] = np.nan
    pitches['Tilt'] = np.nan
    pitches['SpinDir'] = np.nan
    pitches['SpinEff'] = np.nan
    pitches['GyroDegree'] = np.nan
    pitches['Heart'] = np.nan
    pitches['Shadow'] = np.nan
    pitches['Chase'] = np.nan
    pitches['Waste'] = np.nan
    pitches['Zone'] = np.nan
    lol = 0
    for i in range(len(pitches.pitch_type)):
        # x0 = -1 * pitches.x0.iloc[i]
        pitch_type = pitches.pitch_type.iloc[i]
        v0 = pitches.release_speed.iloc[i]
        vx0 = pitches.vx0.iloc[i]
        ax = pitches.ax.iloc[i]
        vy0 = pitches.vy0.iloc[i]
        ay = pitches.ay.iloc[i]
        vz0 = pitches.vz0.iloc[i]
        az = pitches.az.iloc[i]
        pfx_x = pitches.pfx_x.iloc[i]
        pfx_z = pitches.pfx_z.iloc[i]
        plate_x = pitches.plate_x.iloc[i]
        plate_z = pitches.plate_z.iloc[i]
        release_x = pitches.release_pos_x.iloc[i]
        release_y = 60.5-pitches.release_extension.iloc[i]
        release_z = pitches.release_pos_z.iloc[i]
        spin_rate = pitches.release_spin_rate.iloc[i]
        sz_top = pitches.sz_top.iloc[i]
        sz_bot = pitches.sz_bot.iloc[i]
        #p_throws = pitches.p_throws.iloc[i]

        sz = sz_top - sz_bot
        sz_mid = sz_top - sz/2

        top_heart = sz_mid + 0.333 * sz
        bot_heart = sz_mid - 0.333 * sz
        top_shadow = sz_mid + 0.667 * sz
        bot_shadow = sz_mid - 0.667 * sz
        top_chase = sz_mid + 1 * sz
        bot_chase = sz_mid - 1 * sz

        heart = 0
        shadow = 0
        chase = 0
        waste = 0
        zone = 0

        if  abs(plate_x) <= 0.58333 and bot_heart <= plate_z <= top_heart:
            heart = 1
        elif abs(plate_x) <= 1.108333 and bot_shadow <= plate_z <= top_shadow:
            shadow = 1
        elif abs(plate_x) <= 1.66667 and bot_chase <= plate_z <= top_chase:
            chase = 1
        else:
            waste = 1

        if abs(plate_x) <= 0.8333 and sz_bot <= plate_z <= sz_top:
            zone = 1
        else:
             zone = 0

        # time between release and y0 measurement
        t_back_to_release = (-vy0-math.sqrt(vy0**2-2*ay*(50-release_y)))/ay

        # adjust velocity at y0 to be at release
        vx_r = vx0+ax*t_back_to_release
        vy_r = vy0+ay*t_back_to_release
        vz_r = vz0+az*t_back_to_release
        dv0 = v0 - math.sqrt(vx_r**2 + vy_r**2 + vz_r**2)/1.467

        # calculate pitch time also know as tf in Template
        t_c = (-vy_r - math.sqrt(vy_r**2 - 2*ay*(release_y - 17/12))) / ay

        # calcualte x and z movement
        calc_x_mvt = (plate_x-release_x-(vx_r/vy_r)*(17/12-release_y))
        calc_z_mvt = (plate_z-release_z-(vz_r/vy_r)*(17/12-release_y))+0.5*g_fts*t_c**2

        # average velocity
        vx_bar = (2 * vx_r + ax * t_c) / 2
        vy_bar = (2 * vy_r + ay * t_c) / 2
        vz_bar = (2 * vz_r + az * t_c) / 2
        v_bar = math.sqrt(vx_bar**2 + vy_bar**2 + vz_bar**2)

        # drag acceleration
        adrag = -(ax * vx_bar + ay * vy_bar + (az + g_fts) * vz_bar)/v_bar

        # magnus acceleration
        amagx = ax + adrag * vx_bar/v_bar
        amagy = ay + adrag * vy_bar/v_bar
        amagz = az + adrag * vz_bar/v_bar + g_fts
        amag = math.sqrt(amagx**2 + amagy**2 + amagz**2)

        # movement components
        mx = .5 * amagx * (t_c**2)*12
        mz = .5 * amagz * (t_c**2)*12

        # drag/lift coefficients may need work
        Cd = adrag / (v_bar**2 * const)
        Cl = amag / (v_bar**2 * const)

        s = 0.4*Cl/(1-2.32*Cl)
        spin_t = 78.92*s*v_bar

        '''
        # for debugging purposes
        spin_tx = spin_t*(vy_bar*amagz-vz_bar*amagy)/(amag*v_bar)
        spin_ty = spin_t*(vz_bar*amagx-vx_bar*amagz)/(amag*v_bar)
        spin_tz = spin_t*(vx_bar*amagy-vy_bar*amagx)/(amag*v_bar)
        spin_check = math.sqrt(spin_tx**2+spin_ty**2+spin_tz**2)-spin_t
        '''

        # calc spin direction
        phi = 0
        if(amagz > 0):
            phi = math.atan2(amagz, amagx) * 180/math.pi
        else:
            phi = math.atan2(amagz, amagx)*180/math.pi #+360
        spin_dir = phi


        # calc spin eff
        spin_eff = spin_t/spin_rate
        theta = np.nan
        #print(pitch_type, phi, dec_time)
        tMov = 0
        tMov = math.sqrt((-12*pfx_x)**2 + (12*pfx_z)**2)

        pd.set_option('mode.chained_assignment', None)
        pitches.InducedHorzBreak.iloc[i] = -calc_x_mvt
        pitches.InducedVertBreak.iloc[i] = calc_z_mvt
        pitches.SpinDir.iloc[i] = spin_dir
        pitches.GyroDegree.iloc[i] = theta
        pitches.SpinEff.iloc[i] = spin_eff
        pitches.Heart.iloc[i] = heart
        pitches.Shadow.iloc[i] = shadow
        pitches.Chase.iloc[i] = chase
        pitches.Waste.iloc[i] = waste
        pitches.Zone.iloc[i] = zone
        pitches.total_mov.iloc[i] = tMov
    #pitches.to_csv('test.csv')
    return pitches


def get_pitch_types(data):
    pitchcounts = data['pitch_type'].value_counts(dropna=True)
    pitch_types = []
    for i in range(len(pitchcounts)):
        is_pitch = pitchcounts.index[i]
        pitch_types.append(is_pitch)
    return pitch_types


def color_picker(pitch_type):
    color = ''
    if(pitch_type == 'FF'):
        color = '#8C1C13'
    elif(pitch_type == 'FT'):
        color = '#CC5803'
    elif(pitch_type == 'SI'):
        color = '#FE7F2D'
    elif(pitch_type == 'FC'):
        color = '#E08E45'
    elif(pitch_type == 'FS'):
        color = '#F3CA4C'
    elif(pitch_type == 'SL'):
        color = '#274060'
    elif(pitch_type == 'CU'):
        color = '#4EA5D9'
    elif(pitch_type == 'KC'):
        color = '#5BC0EE'
    elif(pitch_type == 'CH'):
        color = '#1446A0'
    elif(pitch_type == 'KN'):
        color = '#712F79'
    elif(pitch_type == 'FO'):
        color = '#03B5AA'
    elif(pitch_type == 'EP'):
        color = '#DBFE97'
    elif(pitch_type == 'SC'):
        color = '#5B9279'
    else:
        color = 'black'
    return color


def plot_release_movement(data):
    pitch_types = get_pitch_types(data)

    fig1 = make_subplots(rows=1, cols=3,horizontal_spacing = 0.03,
            specs=[[{"type": "xy"}, {"type": "xy"}, {"type": "polar"}]],
            subplot_titles=("Release Position", "Pitch Movement", "Spin Direction")
            )

    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        pitch_count = selected_data['pitch_type'].count()
        total_count = data['pitch_type'].count()
        percent = round(pitch_count/total_count*100, 1)
        label = pitch_types[i]

        if(percent < 1.5):
            continue
        else:
            color = color_picker(label)
            c1 = color_picker(label)
            if(label == 'PO' or label == 'IB' or label == 'AB' or label == 'UN'):
                continue
            else:
                fig1.add_trace(
                    go.Scatter(x=selected_data.release_pos_x*-1, y=selected_data.release_pos_z,
                                mode="markers",
                                showlegend=False,
                                line=dict(
                                    color=c1
                                    )
                                ),
                    row=1, col=1
                )

                fig1.add_shape(type="circle",
                                xref="x", yref="y",
                                fillcolor="#742604",
                                x0=-9, y0=-1, x1=9, y1=10/12,
                                line_color="#742604",
                                row=1, col=1
                            )
                fig1.add_shape(
                                dict(type="rect",
                                x0=-9/12, x1=9/12, y0=9/12, y1=11/12,
                                fillcolor = "white",
                                line_width = 0.75,
                                line_color="black"),
                                row=1,col=1,
                            )

                fig1.add_trace(
                    go.Scatter(x=-12*selected_data.pfx_x, y=12*selected_data.pfx_z,
                                mode="markers",
                                name = label,
                                line=dict(
                                    color=c1
                                    )
                                ),
                    row=1, col=2
                )

                fig1.add_trace(
                    go.Scatterpolar(r=selected_data.total_mov, theta=selected_data.spin_axis,
                                mode="markers",
                                showlegend=False,
                                subplot = "polar1",
                                line=dict(
                                    color=c1
                                    )
                                ),
                    row=1, col=3
                )
    fig1.update_annotations(font_size=36)
    fig1.update_layout(
        height=600,
        width=1800,
        font=dict( size=24),
        margin=dict(l=35, r=35, t=40, b=10),
            legend=dict(
                    orientation="h",
                    yanchor="top",
                    y=-0.075,
                    xanchor="center",
                    x=0.5
                ),
            polar1 = dict(
                        radialaxis = dict(range=[0, 30], tickvals = [0,10,20,30]),
                        angularaxis = dict(
                                            rotation = 270,
                                            direction="clockwise",
                                            tickvals = [270, 240, 210, 180, 150,120, 90,60, 30,0, 330, 300],
                                            ticktext = ["3:00","2:00", "1:00", "","11:00", "10:00", "9:00","8:00", "7:00", "6:00", "5:00", "4:00"]
                                            )
                        )
    )

    fig1.update_xaxes(range=[-5, 5], row=1, col=1)
    fig1.update_yaxes(range=[0, 8], row=1, col=1)
    fig1.update_xaxes(range=[-30, 30], row=1, col=2)
    fig1.update_yaxes(range=[-30, 30], row=1, col=2)
    #fig1.show()
    memfile = fig1.write_image("fig1.jpeg")


def plot_location(data):
    pitch_types = get_pitch_types(data)
    gs = gridspec.GridSpec(1, 7)
    fig = plt.figure(figsize=(7.5, 2))
    ax0 = plt.subplot(gs[:, 0])  # FF
    ax1 = plt.subplot(gs[:, 1])  # FT/SI
    ax2 = plt.subplot(gs[:, 2])  # FC
    ax3 = plt.subplot(gs[:, 3])  # SL
    ax4 = plt.subplot(gs[:, 4])  # CU/KC
    ax5 = plt.subplot(gs[:, 5])  # CH
    ax6 = plt.subplot(gs[:, 6])  # FS

    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        label = pitch_types[i]
        pitch_count = selected_data['pitch_type'].count()
        total_count = data['pitch_type'].count()
        percent = round(pitch_count/total_count*100, 1)
        label = pitch_types[i]
        if(percent < 1.5):
            continue
        else:
            sz_x = [.83, .83, -.83, -.83, .83]
            sz_z = [3.5, 1.5, 1.5, 3.5, 3.5]
            xedges, zedges = np.linspace(-2, 2, 20), np.linspace(-0.5, 4.5, 20)
            x = -1*selected_data['plate_x']
            z = selected_data['plate_z']
            hist, xedges, yedges = np.histogram2d(x, z, (xedges, zedges))
            xidx = np.clip(np.digitize(x, xedges), 0, hist.shape[0]-1)
            zidx = np.clip(np.digitize(z, zedges), 0, hist.shape[1]-1)
            c = hist[xidx, zidx]
            if(label == 'FF'):
                ax0.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax0.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FT' or label == 'SI'):
                ax1.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax1.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FC'):
                ax2.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax2.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'SL'):
                ax3.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax3.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'CU' or label == 'KC'):
                ax4.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax4.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'CH'):
                ax5.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax5.plot(sz_x, sz_z, color='black', lw=0.5)
            elif(label == 'FS'):
                ax6.scatter(x, z, c=c, s=1, cmap='YlOrRd')
                ax6.plot(sz_x, sz_z, color='black', lw=0.5)

    ax0.set_xlim(-2, 2)
    ax0.set_ylim(-0.5, 4.5)
    ax0.axis('off')
    ax0.set_title('FF', fontsize=10)
    ax1.set_xlim(-2, 2)
    ax1.set_ylim(-0.5, 4.5)
    ax1.axis('off')
    ax1.set_title('FT/SI', fontsize=10)
    ax2.set_xlim(-2, 2)
    ax2.set_ylim(-0.5, 4.5)
    ax2.axis('off')
    ax2.set_title('FC', fontsize=10)
    ax3.set_xlim(-2, 2)
    ax3.set_ylim(-0.5, 4.5)
    ax3.axis('off')
    ax3.set_title('SL', fontsize=10)
    ax4.set_xlim(-2, 2)
    ax4.set_ylim(-0.5, 4.5)
    ax4.axis('off')
    ax4.set_title('CU/KC', fontsize=10)
    ax5.set_xlim(-2, 2)
    ax5.set_ylim(-0.5, 4.5)
    ax5.axis('off')
    ax5.set_title('CH', fontsize=10)
    ax6.set_xlim(-2, 2)
    ax6.set_ylim(-0.5, 4.5)
    ax6.axis('off')
    ax6.set_title('FS', fontsize=10)
    fig.suptitle('Pitch Locations (Pitcher\'s View)')
    fig.subplots_adjust(top=.72, wspace=0.0, bottom=.0, left=.00, right=1)
    memfile = BytesIO()
    plt.savefig(memfile)
    # plt.show() for debug
    return memfile


# rounds to nearest 15 minutes
def time_round(x, base=1):
    return base * round(x/base)


def convert_to_time(dec_time):
    if(math.isnan(dec_time)):
        time = ''
    else:
        hours = int(dec_time)
        minutes = int((dec_time*60) % 60)
        minutes = time_round(minutes)
        if(minutes == 60):
            hours += 1
            minutes = 0
        if(hours == 0):
            hours = 12
        elif(hours > 12):
            hours = hours - 12
        minutestr = str(minutes)
        if(len(minutestr) < 2):
            minutestr = '0' + minutestr
        time = str(hours) + ':' + minutestr
    return time


def avg_spin_dir(phi, x):
    #phi = phi + 90
    if(x < 0):
        phi = phi/30
        dec_time = phi + 6
    else:
        phi = phi/30
        dec_time = phi - 6

    if (phi < 1):
        dec_time += 12
    elif (phi >= 13):
        dec_time -= 12
    else:
        dec_time += 0

    return dec_time


def remove_outliers(df):
    cols = ['pfx_x', 'pfx_z', 'release_speed', 'release_spin_rate'] # one or more

    Q1 = df[cols].quantile(0.25)
    Q3 = df[cols].quantile(0.75)
    IQR = Q3 - Q1

    data = df[~((df[cols] < (Q1 - 1.5 * IQR)) |(df[cols] > (Q3 + 1.5 * IQR))).any(axis=1)]

    return data


def transform_data(data, spin_data):
    pitch_types = get_pitch_types(data)
    pitches = []
    battedball = []
    for i in range(len(pitch_types)):
        is_pitch = data['pitch_type'] == pitch_types[i]
        selected_data = data[is_pitch]
        count = selected_data['pitch_type'].count()
        if (count > 0):
            total_re = selected_data['delta_run_exp'].sum()
            selected_data = remove_outliers(selected_data)
            label = pitch_types[i]
            pitch_type = ''
            if label == 'KC' or label == 'CS':
                pitch_type = 'CU'
            else:
                pitch_type = label
            # pitch info stuff
            swings_misses = ['swinging_strike', 'swinging_strike_blocked',
                                 'foul_tip', 'swinging_strike_pitchout']
            swm = selected_data[selected_data['description'].isin(swings_misses)].count()[
                'description']
            swings = ['swinging_strike', 'swinging_strike_blocked', 'foul',
                              'foul_tip', 'hit_into_play', 'hit_into_play_no_out',
                              'hit_into_play_score', 'swinging_strike_pitchout']
            total_swings = selected_data.loc[selected_data['description'].isin(swings)].count()[
                'description']
            percentage_used = round(
                (count/data['pitch_type'].count()) * 100, 1)
            avgVelo = round(selected_data['release_speed'].dropna().mean(), 1)
            avgSpinRate = round(selected_data['release_spin_rate'].dropna().mean(), 0)
            if avgSpinRate == np.nan:
                avgSpinRate = 0
            else:
                avgSpinRate = int(avgSpinRate)
            avgHorzBreak = round(-12*selected_data['pfx_x'].dropna().mean(), 1)
            avgVertBreak = round(12*selected_data['pfx_z'].dropna().mean(), 1)
            #bauer_units = (round(avgSpinRate/avgVelo, 0))
            #spin_dir_u = selected_data['SpinDir'].dropna().mean()
            obs_spin_dir = selected_data['spin_axis'].dropna().mean()
            h_break_u = -12*selected_data['pfx_x'].dropna().mean()
            v_break_u = 12*selected_data['pfx_z'].dropna().mean()
            velo_u = selected_data['release_speed'].dropna().mean()
            obs_dec_time = avg_spin_dir(obs_spin_dir, h_break_u)
            obs_tilt = convert_to_time(obs_dec_time)
            inf_spin_dir = math.atan2(v_break_u, h_break_u) * 180/math.pi
            if h_break_u > 0:
                inf_spin_dir = 270 - inf_spin_dir
            else:
                inf_spin_dir = 90 - (inf_spin_dir + 180)

            inf_dec_time = avg_spin_dir(inf_spin_dir, h_break_u)
            inf_tilt = convert_to_time(inf_dec_time)

            obs_string = '2021-1-1 ' + str(obs_tilt) + ':00.0'
            inf_strin = '2021-1-1 ' + str(inf_tilt) + ':00.0'
            date_time_obs = datetime.datetime.strptime(obs_string, '%Y-%m-%d %H:%M:%S.%f')
            date_time_inf = datetime.datetime.strptime(inf_strin, '%Y-%m-%d %H:%M:%S.%f')

            ss = ''
            if date_time_inf > date_time_obs:
                ss = date_time_inf - date_time_obs #make negative
            else:
                ss = date_time_obs - date_time_inf

            seam_shift = int(ss.total_seconds() / 60)
            if date_time_inf > date_time_obs:
                seam_shift = seam_shift*-1

            dir = str(obs_tilt) + "/" + str(inf_tilt) + " (" + str(seam_shift) + ")"

            #obs_dec_time = round(dec_time,3)
            #spin_dir = round(selected_data['spin_axis'].dropna().mean(),0)

            #spin_eff = round(selected_data['SpinEff'].dropna().mean()*100, 1)
            velo_u = selected_data['release_speed'].dropna().mean()

            spin_eff1 = round(selected_data['SpinEff'].dropna().mean(), 3) * 100

            tMov = math.sqrt((avgHorzBreak)**2 + (avgVertBreak)**2)
            spin_rate = selected_data['release_spin_rate'].dropna().mean()
            spin_eff = 0.08731 + 0.004127*velo_u - 0.00014*spin_rate + 0.042529*tMov

            if(spin_eff > 1):
                spin_eff = 1
            else:
                spin_eff += 0

            gyro_degree = int(round(math.degrees(math.acos(spin_eff)),0))
            spin_eff = round(spin_eff * 100, 1)

            #VAA coef
            vaa_int = -15.0738
            rel_speed_int = 0.116686
            rel_height_int = -0.831169
            rel_ext_int = 0.0345948
            az_int = -0.0304108
            ay_int = -0.023317
            vz0_int = 0.0799197
            ind_vert_brk_int = 0.136734
            plate_z_int = 0.854793

            #Drop coef
            drop_int = -139.113
            drop_vert_bk = 1.31986
            drop_az = -0.349845
            drop_vy0 = -0.843272
            drop_rel_speed = -0.24393
            drop_ay = -0.0913047
            drop_ext = 1.36142
            drop_plate_z = -0.221783

            #VAA
            rel_speed = selected_data['release_speed'].dropna().mean() * rel_speed_int
            rel_height = selected_data['release_pos_z'].dropna().mean() * rel_height_int
            rel_ext = selected_data['release_extension'].dropna().mean() * rel_ext_int
            az = selected_data['az'].dropna().mean() * az_int
            ay = selected_data['ay'].dropna().mean() * ay_int
            vx0 = selected_data['vz0'].dropna().mean() * vz0_int
            ind_vert_brk = selected_data['pfx_z'].dropna().mean() * 12 * ind_vert_brk_int
            plate_z = selected_data['plate_z'].dropna().mean() * plate_z_int

            vaa = round(vaa_int + rel_speed + rel_height + rel_ext + az + ay + vx0 + ind_vert_brk + plate_z,2)

            #drop
            d_vb = selected_data['pfx_z'].dropna().mean() * 12 * drop_vert_bk
            d_az = selected_data['az'].dropna().mean() * drop_az
            d_vy0 = selected_data['vy0'].dropna().mean() * drop_vy0
            d_velo = selected_data['release_speed'].dropna().mean() * drop_rel_speed
            d_ay = selected_data['ay'].dropna().mean() * drop_ay
            d_ext = selected_data['release_extension'].dropna().mean() * drop_ext
            d_plate_z = selected_data['plate_z'].dropna().mean() * drop_plate_z

            drop = round(drop_int + d_vb + d_az + d_vy0 + d_velo + d_ay + d_ext + d_plate_z,1)

            # second table
            bbe = selected_data['launch_speed_angle'].dropna().count()
            '''weak = round(selected_data[selected_data['launch_speed_angle'] == 1].count()[
                'launch_speed_angle']/bbe*100, 1)
            topped = round(selected_data[selected_data['launch_speed_angle'] == 2].count()[
                'launch_speed_angle']/bbe*100, 1)
            under = round(selected_data[selected_data['launch_speed_angle'] == 3].count()[
                'launch_speed_angle']/bbe*100, 1)
            flare = round(selected_data[selected_data['launch_speed_angle'] == 4].count()[
                'launch_speed_angle']/bbe*100, 1)
            solid = round(selected_data[selected_data['launch_speed_angle'] == 5].count()[
                'launch_speed_angle']/bbe*100, 1)'''
            barrels = round(selected_data[selected_data['launch_speed_angle'] == 6].count()[
                'launch_speed_angle']/bbe*100, 1)
            heart_pct = round(selected_data[selected_data['Heart'] == 1].count()['Heart']/count*100,1)
            shadow_pct = round(selected_data[selected_data['Shadow'] == 1].count()['Shadow']/count*100,1)
            chase_pct = round(selected_data[selected_data['Chase'] == 1].count()['Chase']/count*100,1)
            waste_pct = round(selected_data[selected_data['Waste'] == 1].count()['Waste']/count*100,1)
            zone_pct = round(selected_data[selected_data['Zone'] == 1].count()['Zone']/count*100,1)
            whiff_rate = round(swm/total_swings*100, 1)
            hardhit = round(selected_data[selected_data['launch_speed'] >= 95].count()[
                'launch_speed_angle']/bbe*100, 1)

            re_100 = round(total_re / count * 100,1)
            re = round(total_re,1)
            run_exp = str(re) + " (" + str(re_100) + ")"

            pitch = [label, pitch_type, percentage_used, avgVelo, avgSpinRate,
                     avgVertBreak, avgHorzBreak,dir]
            bbs = [pitch_type, percentage_used, vaa, drop,
                   heart_pct, shadow_pct,
                   whiff_rate, hardhit, run_exp]
            pitches.append(pitch)
            battedball.append(bbs)
        else:
            continue
    pitches = pd.DataFrame(pitches, columns=['Pitch', 'pitch_type', '% Thrown', 'Velocity',
                                             'Spin Rate', 'vMov',
                                             'hMov', 'Obs/Inf Dir.'])
    pitches = pitches.merge(spin_data, on = 'pitch_type', how = 'left')
    pitches['spin_eff'] = np.nan
    pitches['gyro_deg'] = np.nan
    for i in range(len(pitches.pitch_type)):
        act_spin = pitches.active_spin.iloc[i]
        alan_active = pitches.alan_active_spin_pct.iloc[i]

        if np.isnan(act_spin):
            act_spin = np.nan_to_num(act_spin)
        if act_spin > 1:
            act_spin = 1
        if np.isnan(alan_active):
            alan_active = np.nan_to_num(alan_active)
        if alan_active > 1:
            alan_active = 1

        #print(pitches.pitch_type.iloc[i], ': ', act_spin, ' | ', alan_active)
        spin_eff = str(round(act_spin*  100,1)) + '/' + str(round(alan_active* 100,1))
        act_gyro_degree = int(round(math.degrees(math.acos(act_spin)),0))
        alan_gyro_degree = int(round(math.degrees(math.acos(alan_active)),0))

        gyro_deg = str(act_gyro_degree) + '/' + str(alan_gyro_degree)

        pitches.gyro_deg.iloc[i] = gyro_deg
        pitches.spin_eff.iloc[i] = spin_eff

    pitches = pitches.drop(columns = ['pitch_type','player_id','alan_active_spin_pct', 'active_spin'])
    pitches = pitches.rename(columns={"spin_eff": "Spin Eff.", "gyro_deg": "Gyro Deg."})
    battedball = pd.DataFrame(battedball, columns=['Pitch', '% Thrown', 'VAA',
                                                   'Drop',
                                                   'Heart%', 'Shadow%',
                                                   'Whiff%', 'HardHit%',
                                                   'RV (/100)'])
    pitches = pitches[pitches['% Thrown'] >= 1.5]
    battedball = battedball[battedball['% Thrown'] >= 1.5]
    pitches = pitches.sort_values(by=['% Thrown'], ascending=False)
    battedball = battedball.sort_values(by=['% Thrown'], ascending=False)

    return pitches, battedball


def generate_profile(fname, lname, date1, date2, moves,locs, pitches, battedballs, count, dates):
    document = Document()
    section = document.sections[0]

    margin = 0.5
    section.top_margin = Inches(margin)
    section.bottom_margin = Inches(margin)
    section.left_margin = Inches(margin)
    section.right_margin = Inches(margin)

    paragraph = document.add_paragraph()
    paragraph_format = paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titlestr = 'Pitch Profile for ' + fname + ' ' + lname
    datestr = '\n' + date1 + ' to ' + date2
    run0 = paragraph.add_run(titlestr)
    run1 = paragraph.add_run(datestr)
    font0 = run0.font
    font0.size = Pt(24)
    font0.bold = True
    font0.name = 'Calibri'
    font1 = run1.font
    font1.size = Pt(12)
    font1.name = 'Calibri'

    move_paragraph = document.add_paragraph()
    paragraph_format = move_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = move_paragraph.add_run()
    my_image = run.add_picture("fig1.jpeg", width=Inches(7.5), height=Inches(2.5))
    #my_image = run.add_picture(BytesIO(moves))

    loc_paragraph = document.add_paragraph()
    paragraph_format = loc_paragraph.paragraph_format
    paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = loc_paragraph.add_run()
    my_image = run.add_picture(locs)

    # First row are table headers
    pitchTable = document.add_table(
        pitches.shape[0]+1, pitches.shape[1], style='MediumList2')

    # add the header rows.
    for j in range(pitches.shape[-1]):
        pitchTable.cell(0, j).text = pitches.columns[j]

    # add the rest of the data frame
    for i in range(pitches.shape[0]):
        for j in range(pitches.shape[-1]):
            pitchTable.cell(i+1, j).text = str(pitches.values[i, j])

    break_paragraph = document.add_paragraph('')

    # First row are table headers
    bbTable = document.add_table(
        battedballs.shape[0]+1, battedballs.shape[1], style='MediumList2')

    # add the header rows.
    for j in range(battedballs.shape[-1]):
        bbTable.cell(0, j).text = battedballs.columns[j]

    # add the rest of the data frame
    for i in range(battedballs.shape[0]):
        for j in range(battedballs.shape[-1]):
            bbTable.cell(i+1, j).text = str(battedballs.values[i, j])

    doc_name = ''
    if dates == 'REG':
        doc_name = fname+ "_"+lname+"_REG_SEASON"
    else:
        doc_name = fname+ "_"+lname+"_"+date2

    document.save(doc_name + '.docx')
    #moves.close()
    locs.close()


def main():
    while (1 == 1):
        fname = input('Enter First Name: ')
        lname = input('Enter Last Name: ')
        #team = input('Enter Team Played For (EX: MIN for MN Twins): ')
        date1 = ''
        date2 = ''
        dates = input('Current Season (REG) or Custom Range (CUS)?: ')
        if dates == 'REG':
            date1 = '2021-4-1'
            date2 = '2021-10-3'
        else:
            date1 = input('Enter Start of Date Range (YYYY-MM-DD): ')
            date2 = input('Enter End of Date Range (YYYY-MM-DD): ')
        year = int(date1[:4])
        data, spin_data = import_data(get_number(lname, fname, year), date1, date2, year)
        data = nathan_calculations(data)
        pitchdata, battedballdata = transform_data(data, spin_data)
        releasemovement = plot_release_movement(data)
        locations = plot_location(data)
        pitches = get_pitch_types(data)
        count = 0
        for i in range(len(pitches)):
            count += 1
        generate_profile(fname, lname, date1, date2, releasemovement,
                         locations, pitchdata, battedballdata, count, dates)
        unused_variable = os.system("cls")
        print('Pitch Profile for ' + fname + ' ' + lname + ' Created!')
        dec = input('Enter New Pitcher (Y/N)?: ').upper()
        if(dec == 'N'):
            break
        else:
            continue


if __name__ == '__main__':
    main()
